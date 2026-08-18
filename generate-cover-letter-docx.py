#!/usr/bin/env python3
"""generate-cover-letter-docx.py — Renders a cover letter payload to an editable .docx.

Takes the same JSON payload as generate-cover-letter.mjs (the PDF renderer).
Structure mirrors templates/cover-letter-template.html section for section
so the two outputs stay in sync.

Usage:
  python3 generate-cover-letter-docx.py --payload payload.json
  python3 generate-cover-letter-docx.py --payload payload.json --out output/slug-cover.docx

Letterhead handling (payload.letterhead, all optional):
  - base_docx_path: start from this .docx (its header/footer survive untouched,
    e.g. a real employer-issued letterhead template) instead of a blank document.
  - header_image_path / footer_image_path: used only when base_docx_path is
    absent — inserted as plain in-body pictures at the top/bottom of the page.
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Matches the `[label](https://url)` inline link syntax accepted by
# generate-cover-letter.mjs, so the PDF and DOCX twins render the same links.
LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")

GRAY = RGBColor(0x66, 0x66, 0x66)
LINK_BLUE = RGBColor(0x1A, 0x56, 0xA0)

# Mirrors templates/cover-letter-template.html's CSS (Helvetica/Arial 10pt
# body) so the DOCX twin reads the same as the PDF regardless of what font
# the base template (e.g. a real letterhead .docx) happened to ship with.
BODY_FONT = "Arial"
BODY_SIZE = Pt(10)


def set_body_font(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE


def style_run(run, size=None, bold=None, italic=None, color=None):
    run.font.name = BODY_FONT
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def add_hyperlink(paragraph, url, text):
    # python-docx has no built-in hyperlink support -- build the w:hyperlink
    # run manually via the relationship part, styled to match LINK_BLUE.
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(BODY_SIZE.pt * 2)))
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), str(LINK_BLUE))
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text, bold=None):
    # Splits text on `[label](url)` spans, adding plain styled runs for the
    # surrounding text and real hyperlink runs for each link.
    last = 0
    for m in LINK_RE.finditer(text):
        if m.start() > last:
            style_run(paragraph.add_run(text[last:m.start()]), bold=bold)
        add_hyperlink(paragraph, m.group(2), m.group(1))
        last = m.end()
    if last < len(text):
        style_run(paragraph.add_run(text[last:]), bold=bold)


def add_paragraphs_with_links(doc, text):
    # Mirrors buildParagraphsBlock() in generate-cover-letter.mjs: split on
    # blank lines (or an array) into one Word paragraph per source paragraph.
    if not text:
        return
    paras = text if isinstance(text, list) else re.split(r"\n{2,}", text)
    for para in paras:
        p = doc.add_paragraph()
        add_text_with_links(p, para)


def require(obj, keys, context):
    for key in keys:
        if key not in obj:
            sys.exit(f"ERROR: missing required field: {context}.{key}")


def contact_line_parts(candidate):
    parts = []
    if candidate.get("location"):
        parts.append(candidate["location"])
    if candidate.get("email"):
        parts.append(candidate["email"])
    if candidate.get("phone"):
        parts.append(candidate["phone"])
    if candidate.get("linkedin"):
        parts.append(candidate["linkedin"])
    if candidate.get("github"):
        parts.append(candidate["github"])
    if candidate.get("portfolio_url"):
        parts.append(candidate["portfolio_url"])
    return parts


def add_contact_paragraph(doc, candidate):
    # Single plain-text run, no per-segment hyperlinks -- meant for hand
    # editing, not ATS parsing.
    p = doc.add_paragraph()
    run = p.add_run("  |  ".join(contact_line_parts(candidate)))
    style_run(run, size=Pt(9), color=GRAY)
    return p


def has_style(doc, name):
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def add_achievement(doc, lead, impact, bullet_style_available):
    # Not every base template (e.g. a stripped-down letterhead) ships the
    # built-in "List Bullet" style -- fall back to a literal bullet glyph.
    # (Checked once upfront: python-docx's add_paragraph() appends the
    # paragraph to the document *before* validating the style name, so
    # catching the KeyError after the fact leaves an orphaned empty
    # paragraph behind -- must not attempt the bad style at all.)
    if bullet_style_available:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
        style_run(p.add_run("• "))
    add_text_with_links(p, lead, bold=True)
    add_text_with_links(p, f" {impact}")
    return p


def build_docx(payload, out_path):
    require(payload, ["candidate", "letter"], "payload")
    candidate = payload["candidate"]
    letter = payload["letter"]
    require(candidate, ["name"], "candidate")
    require(letter, ["role_title", "opening", "profile_intro"], "letter")

    letterhead = payload.get("letterhead") or {}
    base_docx_path = letterhead.get("base_docx_path")

    if base_docx_path:
        base = Path(base_docx_path)
        if not base.exists():
            sys.exit(f"ERROR: letterhead base_docx_path not found: {base}")
        doc = Document(str(base))
        set_body_font(doc)
        # The template ships with a single empty placeholder paragraph in the
        # body -- header/footer sections are untouched, only body content
        # below is ours to fill in.
        for p in list(doc.paragraphs):
            el = p._element
            el.getparent().remove(el)
    else:
        doc = Document()
        set_body_font(doc)
        header_image = letterhead.get("header_image_path")
        if header_image and Path(header_image).exists():
            doc.add_picture(header_image, width=Inches(6.5))

        name_p = doc.add_paragraph()
        name_run = name_p.add_run(candidate["name"])
        style_run(name_run, size=Pt(14), bold=True)

        add_contact_paragraph(doc, candidate)

    # --- Role title + dateline ---
    role_p = doc.add_paragraph()
    role_run = role_p.add_run(f"Cover Letter: {letter['role_title']}")
    style_run(role_run, size=Pt(11), bold=True)

    dateline_bits = [b for b in [letter.get("company"), letter.get("city"), letter.get("date")] if b]
    if dateline_bits:
        dl_p = doc.add_paragraph()
        dl_run = dl_p.add_run("    ".join(dateline_bits))
        style_run(dl_run, size=Pt(9), color=GRAY)

    doc.add_paragraph()  # spacer

    # --- Greeting ---
    if letter.get("greeting"):
        doc.add_paragraph(letter["greeting"])

    # --- Opening / profile intro ---
    opening_p = doc.add_paragraph()
    add_text_with_links(opening_p, letter["opening"])
    intro_p = doc.add_paragraph()
    add_text_with_links(intro_p, letter["profile_intro"])

    # --- Achievements ---
    # Optional `group` label (e.g. "At CHOP, I:") renders as a bold heading
    # before its own run of bullets whenever it differs from the previous
    # achievement's group; achievements without a `group` field render as a
    # flat bullet list, unchanged from prior behavior.
    bullet_style_available = has_style(doc, "List Bullet")
    current_group = object()  # sentinel, never equals a real group value or None
    for ach in letter.get("achievements") or []:
        group = ach.get("group")
        if group and group != current_group:
            heading_p = doc.add_paragraph()
            style_run(heading_p.add_run(group), bold=True)
        current_group = group
        add_achievement(doc, ach.get("lead", ""), ach.get("impact", ""), bullet_style_available)

    # --- Problems section ---
    add_paragraphs_with_links(doc, letter.get("problems_section"))

    # --- Closing ---
    add_paragraphs_with_links(doc, letter.get("closing"))

    if letter.get("language_closing"):
        lc_p = doc.add_paragraph()
        lc_run = lc_p.add_run(letter["language_closing"])
        style_run(lc_run, italic=True)

    # --- Footnotes ---
    for fn in letter.get("footnotes") or []:
        if isinstance(fn, dict):
            text = " ".join(x for x in [fn.get("marker"), fn.get("text"), fn.get("url")] if x)
        else:
            text = str(fn)
        fn_p = doc.add_paragraph()
        fn_run = fn_p.add_run(text)
        style_run(fn_run, size=Pt(8), color=GRAY)

    # --- Letterhead footer image (only in the no-base-docx path) ---
    if not base_docx_path:
        footer_image = letterhead.get("footer_image_path")
        if footer_image and Path(footer_image).exists():
            doc.add_paragraph()
            doc.add_picture(footer_image, width=Inches(6.5))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--payload", required=True)
    parser.add_argument("--out")
    args = parser.parse_args()

    payload_path = Path(args.payload)
    if not payload_path.exists():
        sys.exit(f"ERROR: payload file not found: {payload_path}")

    payload = json.loads(payload_path.read_text())

    if args.out:
        out_path = Path(args.out)
    elif payload.get("output_path"):
        out_path = Path(payload["output_path"]).with_suffix(".docx")
    else:
        company = (payload.get("letter", {}).get("company") or "company").lower().replace(" ", "-")
        role = (payload.get("letter", {}).get("role_title") or "role").lower().replace(" ", "-")[:30]
        out_path = Path("output") / f"{company}-{role}-cover.docx"

    build_docx(payload, out_path)
    print(f"Cover letter DOCX: {out_path}")


if __name__ == "__main__":
    main()
